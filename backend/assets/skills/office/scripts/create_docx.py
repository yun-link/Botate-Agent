#!/usr/bin/env python3
"""
create_docx.py - 从JSON创建Word文档，与read_docx.py格式完全兼容
Usage: python create_docx.py input.json [output.docx]
"""

import sys
import json
import re
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, font_info):
    """设置文本运行的字体"""
    if not font_info:
        return
    
    font = run.font
    
    # 基本样式
    if font_info.get('bold') is not None:
        font.bold = font_info['bold']
    if font_info.get('italic') is not None:
        font.italic = font_info['italic']
    if font_info.get('underline') is not None:
        font.underline = font_info['underline']
    if font_info.get('strike') is not None:
        font.strike = font_info['strike']
    
    # 字体大小
    if font_info.get('font_size'):
        font.size = Pt(font_info['font_size'])
    
    # 字体颜色
    if font_info.get('font_color'):
        try:
            color = font_info['font_color']
            if isinstance(color, str) and len(color) >= 6:
                # 处理RGB格式如"FF0000"
                rgb = color[-6:]  # 取最后6位
                font.color.rgb = RGBColor(
                    int(rgb[0:2], 16),
                    int(rgb[2:4], 16),
                    int(rgb[4:6], 16)
                )
        except Exception:
            pass
    
    # 字体名称（通过XML设置以确保中文生效）
    r = run._element
    rPr = r.get_or_add_rPr()
    
    # 设置中文字体
    if font_info.get('font_name_eastasia'):
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_info['font_name_eastasia'])
    
    # 设置西文字体
    if font_info.get('font_name'):
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_info['font_name'])
        rFonts.set(qn('w:hAnsi'), font_info['font_name'])
        font.name = font_info['font_name']
    
    # 高亮颜色
    if font_info.get('highlight_color') and font_info['highlight_color'] != 'None':
        try:
            font.highlight_color = int(font_info['highlight_color'])
        except (ValueError, TypeError):
            pass


def set_paragraph_alignment(paragraph, alignment_name):
    """设置段落对齐方式"""
    mapping = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if alignment_name in mapping:
        paragraph.alignment = mapping[alignment_name]


def set_paragraph_list(paragraph, list_info):
    """设置段落列表/编号样式"""
    if not list_info:
        return
    
    # 获取或创建pPr
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    
    # 创建numPr
    numPr = OxmlElement('w:numPr')
    
    # 设置ilvl（层级）
    ilvl = list_info.get('level', 0)
    ilvl_elem = OxmlElement('w:ilvl')
    ilvl_elem.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_elem)
    
    # 设置numId（编号ID）
    num_id = list_info.get('num_id', 1)
    if num_id:
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), str(num_id))
        numPr.append(numId_elem)
    
    pPr.append(numPr)


def add_formatted_paragraph(doc, para_data):
    """添加格式化的段落"""
    # 创建段落
    paragraph = doc.add_paragraph()
    
    # 设置段落样式
    if para_data.get('style'):
        try:
            paragraph.style = para_data['style']
        except KeyError:
            # 如果样式不存在，使用默认样式
            pass
    
    # 设置对齐
    if para_data.get('alignment'):
        set_paragraph_alignment(paragraph, para_data['alignment'])
    
    # 设置列表
    if para_data.get('list_info'):
        set_paragraph_list(paragraph, para_data['list_info'])
    
    # 添加文本runs
    runs_data = para_data.get('runs', [])
    
    if not runs_data:
        # 如果没有runs数据但有text，创建一个默认run
        if para_data.get('text'):
            run = paragraph.add_run(para_data['text'])
    else:
        for run_info in runs_data:
            text = run_info.get('text', '')
            if text or run_info.get('preserve_empty', False):
                run = paragraph.add_run(text)
                set_run_font(run, run_info)
    
    return paragraph


def add_table(doc, table_data):
    """添加表格"""
    rows = table_data.get('rows', 0)
    cols = table_data.get('columns', 0)
    
    if rows == 0 or cols == 0:
        return None
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'  # 默认样式
    
    data = table_data.get('data', [])
    
    for row_idx, row_data in enumerate(data):
        if row_idx >= rows:
            break
        
        row = table.rows[row_idx]
        
        for col_idx, cell_data in enumerate(row_data):
            if col_idx >= cols:
                break
            
            if cell_data is None:
                continue
            
            cell = row.cells[col_idx]
            
            # 清空默认段落
            cell.text = ''
            
            # 添加段落内容
            paragraphs = cell_data.get('paragraphs', [])
            if not paragraphs and cell_data.get('text'):
                # 兼容简单文本格式
                p = cell.add_paragraph(cell_data.get('text'))
            else:
                for para in paragraphs:
                    # 如果cell已有默认段落，复用它
                    if cell.paragraphs and cell.paragraphs[0].text == '':
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    
                    # 复制段落内容
                    if para.get('runs'):
                        for run_info in para.get('runs', []):
                            run = p.add_run(run_info.get('text', ''))
                            set_run_font(run, run_info)
                    elif para.get('text'):
                        p.add_run(para.get('text'))
                    
                    # 对齐方式
                    if para.get('alignment'):
                        set_paragraph_alignment(p, para.get('alignment'))
            
            # 设置列宽
            if cell_data.get('width'):
                try:
                    cell.width = Inches(cell_data['width'])
                except Exception:
                    pass
    
    # 处理合并单元格
    merge_cells(table, data)
    
    return table


def merge_cells(table, data):
    """处理表格中的合并单元格"""
    # 收集合并信息
    merge_map = {}  # (row, col) -> merge_info
    
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            if not cell_data:
                continue
            
            # 检查grid_span（横向合并）
            grid_span = cell_data.get('grid_span')
            if grid_span and grid_span > 1:
                merge_map[(row_idx, col_idx)] = {
                    'type': 'horizontal',
                    'span': grid_span
                }
            
            # 检查v_merge（纵向合并）
            v_merge = cell_data.get('v_merge')
            if v_merge:
                if v_merge == 'restart':
                    merge_map[(row_idx, col_idx)] = {
                        'type': 'vertical',
                        'action': 'restart'
                    }
                elif v_merge == 'continue':
                    merge_map[(row_idx, col_idx)] = {
                        'type': 'vertical',
                        'action': 'continue'
                    }
    
    # 执行合并
    # 注意：openpyxl风格的v_merge标记需要转换为python-docx的合并操作
    
    # 先处理横向合并
    for (row_idx, col_idx), info in merge_map.items():
        if info['type'] == 'horizontal':
            span = info['span']
            try:
                start_cell = table.cell(row_idx, col_idx)
                end_cell = table.cell(row_idx, col_idx + span - 1)
                start_cell.merge(end_cell)
            except Exception:
                pass
    
    # 重新获取表格结构（因为横向合并改变了结构）
    # 处理纵向合并需要更谨慎
    for (row_idx, col_idx), info in list(merge_map.items()):
        if info['type'] == 'vertical' and info['action'] == 'restart':
            # 找到连续的continue范围
            end_row = row_idx
            for r in range(row_idx + 1, len(data)):
                cell_data = data[r][col_idx] if col_idx < len(data[r]) else None
                if cell_data and cell_data.get('v_merge') == 'continue':
                    end_row = r
                else:
                    break
            
            if end_row > row_idx:
                try:
                    start_cell = table.cell(row_idx, col_idx)
                    end_cell = table.cell(end_row, col_idx)
                    start_cell.merge(end_cell)
                except Exception:
                    pass


def set_document_metadata(doc, metadata):
    """设置文档元数据"""
    core_props = doc.core_properties
    
    if metadata.get('title'):
        core_props.title = metadata['title']
    if metadata.get('author'):
        core_props.author = metadata['author']


def create_docx_from_json(json_path, output_path=None):
    """从JSON创建Word文档"""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 创建文档
    doc = Document()
    
    # 设置元数据
    if data.get('metadata'):
        set_document_metadata(doc, data['metadata'])
    
    # 处理内容
    content = data.get('content', [])
    
    for item in content:
        item_type = item.get('type')
        
        if item_type == 'paragraph':
            # 跳过完全空白的段落（可选，保留以保持一致性）
            add_formatted_paragraph(doc, item)
        
        elif item_type == 'table':
            add_table(doc, item)
    
    # 保存文档
    output = output_path or json_path.replace('.json', '.docx')
    doc.save(output)
    return output


def verify_round_trip(original_docx, json_path, regenerated_docx):
    """
    验证双向一致性（简单对比）
    注意：由于Word内部表示的复杂性，完全字节级一致很难保证，
    但内容和结构应该保持一致
    """
    from docx import Document as DocReader
    
    orig = DocReader(original_docx)
    regen = DocReader(regenerated_docx)
    
    issues = []
    
    # 对比段落数
    if len(orig.paragraphs) != len(regen.paragraphs):
        issues.append(f"段落数不一致: 原{len(orig.paragraphs)} vs 新{len(regen.paragraphs)}")
    
    # 对比表格数
    if len(orig.tables) != len(regen.tables):
        issues.append(f"表格数不一致: 原{len(orig.tables)} vs 新{len(regen.tables)}")
    
    # 对比文本内容（简化）
    orig_text = '\n'.join([p.text for p in orig.paragraphs])
    regen_text = '\n'.join([p.text for p in regen.paragraphs])
    
    if orig_text != regen_text:
        # 找出差异
        if len(orig_text) != len(regen_text):
            issues.append(f"文本长度不一致: 原{len(orig_text)} vs 新{len(regen_text)}")
    
    return len(issues) == 0, issues


def main():
    parser = argparse.ArgumentParser(
        description='从 JSON 创建 Word 文档',
        usage='python create_docx.py input.json [output.docx]'
    )
    parser.add_argument(
        'input',
        help='输入的 JSON 文件路径'
    )
    parser.add_argument(
        'output',
        nargs='?',
        default=None,
        help='输出的 DOCX 文件路径（可选，默认为 input.json 对应的.docx 文件）'
    )
    
    args = parser.parse_args()
    
    json_path = args.input
    output_path = args.output
    
    try:
        result_path = create_docx_from_json(json_path, output_path)
        print(f"✓ 创建成功: {result_path}")
        
        # 显示统计
        doc = Document(result_path)
        print(f"  段落数: {len(doc.paragraphs)}")
        print(f"  表格数: {len(doc.tables)}")
        
    except FileNotFoundError:
        print(f"✗ 错误: 文件不存在 {json_path}")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"✗ 错误: JSON格式无效 {json_path}")
        sys.exit(1)
    except Exception as e:
        print(f"✗ 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
