#!/usr/bin/env python3
"""
docx2json.py - 将Word文档转换为JSON格式，保留完整格式特征
Usage: python docx2json.py input.docx [output.json]
"""

import sys
import json
import re
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def get_alignment_name(alignment):
    """获取段落对齐方式名称"""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: 'left',
        WD_ALIGN_PARAGRAPH.CENTER: 'center',
        WD_ALIGN_PARAGRAPH.RIGHT: 'right',
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
    }
    return mapping.get(alignment, 'left')


def get_run_formatting(run):
    """提取文本运行的格式特征"""
    font = run.font
    
    # 获取字体颜色
    color = None
    if font.color and font.color.rgb:
        color = str(font.color.rgb)
    
    # 获取字体大小（转换为磅值）
    size = None
    if font.size:
        size = font.size.pt
    
    # 获取中文字体和西文字体
    eastasia_font = None
    ascii_font = None
    if run._element.rPr is not None:
        rFonts = run._element.rPr.rFonts
        if rFonts is not None:
            eastasia_font = rFonts.get(qn('w:eastAsia'))
            ascii_font = rFonts.get(qn('w:ascii'))
    
    return {
        'text': run.text,
        'bold': bool(font.bold),
        'italic': bool(font.italic),
        'underline': bool(font.underline),
        'strike': bool(font.strike),
        'font_size': size,
        'font_color': color,
        'font_name': ascii_font or font.name,
        'font_name_eastasia': eastasia_font,
        'highlight_color': font.highlight_color if hasattr(font, 'highlight_color') else None
    }


def parse_paragraph(paragraph):
    """解析段落，保留完整格式"""
    # 获取段落样式
    style_name = paragraph.style.name if paragraph.style else None
    
    # 解析每个文本运行的格式
    runs = [get_run_formatting(run) for run in paragraph.runs if run.text]
    
    # 合并纯文本（用于快速预览）
    full_text = ''.join(run['text'] for run in runs)
    
    # 检测列表类型
    list_info = None
    if paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.numPr
        if numPr is not None:
            ilvl = numPr.ilvl
            numId = numPr.numId
            list_info = {
                'level': ilvl.val if ilvl else 0,
                'num_id': numId.val if numId else None
            }
    
    return {
        'type': 'paragraph',
        'text': full_text,
        'style': style_name,
        'alignment': get_alignment_name(paragraph.alignment),
        'runs': runs,
        'list_info': list_info,
        'is_empty': not full_text.strip()
    }


def parse_table(table):
    """解析表格结构"""
    rows = []
    
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # 递归解析单元格内的段落
            cell_paragraphs = [parse_paragraph(p) for p in cell.paragraphs]
            
            # 获取单元格合并信息
            cell_info = {
                'paragraphs': cell_paragraphs,
                'text': cell.text,
                'width': cell.width.inches if cell.width else None
            }
            
            # 检测合并单元格（通过XML属性）
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is not None:
                gridSpan = tcPr.gridSpan
                if gridSpan is not None:
                    cell_info['grid_span'] = gridSpan.val
                
                # 垂直合并检测
                vMerge = tcPr.vMerge
                if vMerge is not None:
                    cell_info['v_merge'] = vMerge.val if vMerge.val else 'continue'
            
            cells.append(cell_info)
        rows.append(cells)
    
    return {
        'type': 'table',
        'rows': len(table.rows),
        'columns': len(table.columns),
        'data': rows
    }


def parse_document(doc_path):
    """解析整个文档"""
    doc = Document(doc_path)
    
    # 提取文档属性
    core_props = doc.core_properties
    metadata = {
        'title': core_props.title,
        'author': core_props.author,
        'created': str(core_props.created) if core_props.created else None,
        'modified': str(core_props.modified) if core_props.modified else None,
        'paragraphs_count': len(doc.paragraphs),
        'tables_count': len(doc.tables)
    }
    
    # 按文档顺序解析内容（段落和表格混合）
    content = []
    
    # 使用element遍历保持原始顺序
    for element in doc.element.body:
        tag = element.tag
        
        if tag.endswith('p'):  # 段落
            # 找到对应的paragraph对象
            for para in doc.paragraphs:
                if para._p is element:
                    content.append(parse_paragraph(para))
                    break
                    
        elif tag.endswith('tbl'):  # 表格
            # 找到对应的table对象
            for table in doc.tables:
                if table._tbl is element:
                    content.append(parse_table(table))
                    break
    
    return {
        'metadata': metadata,
        'content': content
    }


def main():
    parser = argparse.ArgumentParser(
        description='将 Word 文档转换为 JSON 格式',
        usage='python docx2json.py input.docx [output.json]'
    )
    parser.add_argument(
        'input',
        help='输入的 DOCX 文件路径'
    )
    parser.add_argument(
        'output',
        nargs='?',
        default=None,
        help='输出的 JSON 文件路径（可选，默认为 input.docx 对应的.json 文件）'
    )
    
    args = parser.parse_args()
    
    input_path = args.input
    output_path = args.output or input_path.replace('.docx', '.json')
    
    try:
        result = parse_document(input_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        print(f"✓ 转换成功: {input_path} -> {output_path}")
        print(f"  段落数: {result['metadata']['paragraphs_count']}")
        print(f"  表格数: {result['metadata']['tables_count']}")
        
    except FileNotFoundError:
        print(f"✗ 错误: 文件不存在 {input_path}")
        sys.exit(1)
    except Exception as e:
        print(f"✗ 错误: {str(e)}")
        sys.exit(1)


if __name__ == '__main__':
    main()
